"""Generate the Iteration 1 report in an IEEE-inspired academic-paper style."""

# ruff: noqa: E501 -- Long report paragraphs are clearer as complete source strings.

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = (
    PROJECT_ROOT / "docs" / "Iteration_1.docx",
    PROJECT_ROOT / "docs" / "Smart_Livestock_Gate_Iteration_1_Report.docx",
)
TRAINING_RESULTS = PROJECT_ROOT / "artifacts" / "training" / "cattle_detector"
TEST_RESULTS = PROJECT_ROOT / "artifacts" / "training" / "cattle_detector_test"


def set_run_font(run, name: str = "Times New Roman", size: float = 10) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_columns(section, count: int, space_twips: int = 360) -> None:
    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section_properties.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), str(space_twips))


def configure_section(section, columns: int) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    set_columns(section, columns)


def configure_document(document: Document) -> None:
    configure_section(document.sections[0], 1)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)


def add_body(document: Document, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.16)
    set_run_font(paragraph.add_run(text))


def add_section_heading(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"{number}. {title.upper()}")
    set_run_font(run, size=10)
    run.bold = True


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    set_run_font(run)
    run.italic = True


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def add_table_caption(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"TABLE {number}\n{title.upper()}")
    set_run_font(run, size=8)


def add_table(document: Document, number: str, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_table_caption(document, number, title)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, heading in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(heading)
        set_run_font(run, size=8)
        run.bold = True
    prevent_row_split(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(values):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=8)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(document: Document, path: Path, number: str, caption: str, width: float) -> None:
    if not path.is_file():
        return
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.paragraph_format.space_before = Pt(5)
    image_paragraph.paragraph_format.space_after = Pt(1)
    image_paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption_paragraph.paragraph_format.space_after = Pt(3)
    run = caption_paragraph.add_run(f"Fig. {number}. {caption}")
    set_run_font(run, size=8)


def add_reference(document: Document, number: int, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(2)
    set_run_font(paragraph.add_run(f"[{number}] {text}"), size=8)


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("Smart Livestock Gate: Cattle Detection from a Top-Down Camera")
    set_run_font(title_run, size=19)
    title_run.bold = True

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(1)
    set_run_font(author.add_run("CMPT 310 Project Team"), size=11)

    affiliation = document.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(8)
    set_run_font(affiliation.add_run("School of Computing Science, Simon Fraser University\nIteration 1 — July 24, 2026"), size=9)

    abstract = document.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.left_indent = Inches(0.45)
    abstract.paragraph_format.right_indent = Inches(0.45)
    abstract.paragraph_format.space_after = Pt(4)
    label = abstract.add_run("Abstract—")
    set_run_font(label, size=9)
    label.bold = True
    text = abstract.add_run(
        "This paper reports the first working milestone of Smart Livestock Gate, a software system that detects cattle in overhead imagery and presents the result in a browser. We prepared the CattleEyeView data for supervised object detection, fine-tuned a compact YOLO11n network, added a Flask inference API, and built a React interface that compares model predictions with human annotations. The current detector reached 94.48% precision, 86.75% recall, and 73.88% mAP averaged from IoU 0.50 to 0.95. These results are encouraging, but they are preliminary: a later audit found 432 duplicated frames across validation and test. No test frame was used for gradient training, although the overlap weakens the independence of checkpoint selection and final evaluation. The next iteration will rebuild the split at the video level and report corrected results."
    )
    set_run_font(text, size=9)

    keywords = document.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.left_indent = Inches(0.45)
    keywords.paragraph_format.right_indent = Inches(0.45)
    keywords.paragraph_format.space_after = Pt(6)
    label = keywords.add_run("Index Terms—")
    set_run_font(label, size=9)
    label.bold = True
    set_run_font(keywords.add_run("cattle detection, computer vision, object detection, transfer learning, precision livestock farming"), size=9)


def build_report() -> Document:
    document = Document()
    configure_document(document)
    add_title_block(document)

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, 2)

    add_section_heading(document, "I", "Introduction")
    add_body(
        document,
        "The goal of Smart Livestock Gate is to turn a fixed overhead camera into a useful livestock-monitoring tool. A frame from the camera is the input. In this iteration, the output is a set of cattle bounding boxes, confidence scores, and a count for that frame. Later iterations can connect detections over time and count animals only when a tracked individual crosses a virtual gate. This gives the project a clear path from perception to tracking and event-level counting.",
        indent=False,
    )
    add_body(
        document,
        "CattleEyeView was chosen because it matches that setting closely. The dataset contains 14 top-down videos, 30,703 extracted frames, and annotations for detection, tracking, pose, segmentation, and crossing events [1]. The first milestone deliberately focuses on one problem—cattle detection—so that the learning pipeline can be tested end to end before temporal logic is added.",
    )

    add_section_heading(document, "II", "Course Requirement and Pretrained Weights")
    add_body(
        document,
        "The course project description says that the core learning component must be implemented or customized by the team and that the project cannot consist only of calling a pretrained model. Our system does use pretrained YOLO11n weights as its starting point. That choice is transfer learning, not off-the-shelf inference. The distinction matters: an unchanged pretrained detector would simply be run on the images, while our detector is trained again using CattleEyeView labels and evaluated as a new one-class model.",
        indent=False,
    )
    add_body(
        document,
        "During fine-tuning, the output is configured for one class, cattle, and almost the entire network remains trainable. The training log records 2,590,035 parameters and 2,590,019 parameters receiving gradients; only a small distribution-focal-loss convolution was frozen by the framework. The resulting checkpoint therefore reflects supervised learning from the project data rather than a fixed collection of borrowed predictions. Ultralytics documents pretrained initialization as one supported starting point for custom training [2], and PyTorch describes the same procedure as fine-tuning a pretrained network for a new task [3].",
    )
    add_body(
        document,
        "This approach fits the assignment because the team prepared the labels, selected the input-output task, configured and trained the learning system, recorded its settings, evaluated its predictions, and integrated it into working software. We do not claim to have invented the YOLO architecture. For the final paper, we will make the amount of learning even clearer by comparing the fine-tuned checkpoint with an untouched pretrained baseline on the same clean test set. A randomly initialized run may also be included if time and hardware permit.",
    )

    add_section_heading(document, "III", "Dataset Preparation")
    add_body(
        document,
        "The downloaded material is stored under the project directory, so the program does not depend on a separate machine-specific dataset root. Raw data remains outside Git because the image archive is about 11 GB, but the code, configuration, and instructions needed to reproduce the layout are versioned. The preparation command validates filenames and annotations, then creates hardlinks in a YOLO-compatible directory. Hardlinks avoid making a second full copy of the images.",
        indent=False,
    )
    add_body(
        document,
        "Dataset inspection also exposed 24,051 COCO records whose image identifiers were inconsistent even though their filenames were usable. The importer resolves those records by filename and reports the normalization instead of silently discarding them. The prepared detector layout currently contains 2,714 training images, 1,136 validation images, and 977 test images.",
    )

    add_section_heading(document, "IV", "Method and Training Protocol")
    add_body(
        document,
        "We fine-tuned YOLO11n as a single-class detector. Each image is resized to 512 by 512 pixels and passed through the network. The network predicts candidate boxes and confidence scores; low-confidence candidates are removed, and non-maximum suppression reduces duplicate boxes around the same animal. Training ran for ten epochs with deterministic behavior enabled and seed 42. The small input and batch size were selected to fit the available laptop GPU.",
        indent=False,
    )
    add_table(
        document,
        "I",
        "Training Configuration",
        ["Setting", "Value"],
        [
            ["Model", "YOLO11n, one cattle class"],
            ["Initialization", "Pretrained weights, then fine-tuned"],
            ["Epochs / image size", "10 / 512 × 512"],
            ["Batch / seed", "4 / 42, deterministic"],
            ["Hardware", "NVIDIA RTX 2050, 4 GB"],
            ["Software", "Python 3.13.2; PyTorch 2.12.1; Ultralytics 8.4.105"],
            ["Elapsed time", "1,302 s, including evaluation"],
        ],
    )
    add_body(
        document,
        "The command-line workflow checks the prepared dataset, selects CUDA when it is available, trains the model, retains the best validation checkpoint, evaluates that checkpoint, and writes a machine-readable report. This makes the experiment repeatable without relying on notebook state or manual file selection.",
    )

    add_section_heading(document, "V", "Preliminary Results")
    add_body(
        document,
        "Table II summarizes the first test run. Precision was higher than recall, which means false alarms were less common than missed animals. The stricter mAP result shows that localization quality falls as the required overlap with the human box increases. On the five examples shown in the interface, four predicted the annotated number of animals. The remaining example missed a partially visible animal at the edge of the frame. Keeping that example in the demonstration is useful because it makes the detector's limitation visible rather than presenting only its easiest successes.",
        indent=False,
    )
    add_table(
        document,
        "II",
        "Preliminary Detection Results",
        ["Metric", "Result"],
        [
            ["Precision", "94.48%"],
            ["Recall", "86.75%"],
            ["mAP@0.50", "94.03%"],
            ["mAP@0.75", "82.38%"],
            ["mAP@0.50:0.95", "73.88%"],
            ["Warm example inference", "approximately 88–94 ms"],
        ],
    )

    # Use the full page width for the dense ten-panel training plot.
    wide_figure_section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(wide_figure_section, 1)
    add_figure(
        document,
        TRAINING_RESULTS / "results.png",
        "1",
        "Training and validation losses and detection metrics over ten epochs. These curves describe the completed run but do not remove the split limitation discussed below.",
        6.65,
    )
    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, 2)

    add_section_heading(document, "VI", "Evaluation Limitation")
    add_body(
        document,
        "A frame-level audit performed after training found no exact overlap between training and validation and no overlap between training and test. It did find 432 exact frames shared by validation and test. Both partitions draw from videos 01, 05, 07, 09, and 10. The issue is not direct training leakage—the shared test images were never used for gradient updates—but it is model-selection leakage because validation results influenced which epoch was saved as the best checkpoint.",
        indent=False,
    )
    add_table(
        document,
        "III",
        "Exact Frame Overlap Audit",
        ["Partitions", "Shared frames"],
        [
            ["Train–validation", "0"],
            ["Train–test", "0"],
            ["Validation–test", "432"],
        ],
    )
    add_body(
        document,
        "For that reason, the numbers in Table II must be treated as preliminary rather than as an unbiased final score. The corrected experiment will keep the five official test videos untouched and create a new validation set from the nine training videos, with entire videos assigned to only one partition. The overlap audit will then be rerun before retraining.",
    )
    add_figure(
        document,
        TEST_RESULTS / "confusion_matrix_normalized.png",
        "2",
        "Normalized confusion matrix from the preliminary evaluation. The figure is retained as an iteration artifact, not as the final test result.",
        3.0,
    )

    add_section_heading(document, "VII", "Software Demonstration")
    add_body(
        document,
        "The trained checkpoint is served through a Flask API and used by the React frontend. The backend reports project readiness, returns a fixed set of safe example identifiers, serves the selected images, and performs inference. It does not expose arbitrary local paths. A prediction response contains normalized box coordinates, labels, confidence scores, a count, the image dimensions, the confidence threshold, and inference time.",
        indent=False,
    )
    add_body(
        document,
        "The browser interface places green model boxes and amber human annotations on the same image. It also reports predicted and annotated counts and briefly explains image preparation, feature extraction, box prediction, confidence filtering, and duplicate removal. This turns the frontend into more than a polished output screen: a viewer can inspect what the model believed and where it failed.",
    )

    add_section_heading(document, "VIII", "Verification and Lessons")
    add_body(
        document,
        "The iteration passed 13 Python tests, Ruff static analysis, frontend linting, a production Vite build, repository whitespace checks, and live HTTP checks for the frontend and API. The most practical training problem was that a normal package installation selected CPU-only PyTorch. Installing the official CUDA build allowed the RTX 2050 to be used, although its 4 GB memory still required conservative settings.",
        indent=False,
    )
    add_body(
        document,
        "The larger lesson came from evaluation. Strong metrics are not enough if the provenance of each partition has not been checked. Auditing exact frames and source videos should have happened before training, not afterward. Recording the mistake in this report is important because a credible AI project should explain where its evidence is weaker as well as where the model performs well.",
    )

    add_section_heading(document, "IX", "Next Iteration")
    add_body(
        document,
        "The immediate task is to rebuild the split, verify zero video and frame overlap, and retrain the detector under the same documented settings. We will then compare the customized detector with an untouched pretrained baseline. Once clean detector results are available, the next technical milestone is multi-object tracking with stable identities, followed by a virtual gate that produces one directional crossing event per track. Tracking will be evaluated with identity switches and fragmentation, while gate counting will be compared with the 763 crossing annotations supplied by CattleEyeView.",
        indent=False,
    )

    add_section_heading(document, "X", "Conclusion")
    add_body(
        document,
        "Iteration 1 produced a complete supervised-learning demonstration: local dataset preparation, project-specific fine-tuning, quantitative evaluation, an inference API, and a browser interface. The system already makes understandable cattle detections on real overhead images. Its current scores are promising, but the validation–test overlap prevents them from serving as the final claim. Correcting that split and adding a baseline comparison will make the next iteration both technically stronger and easier to defend as original course work.",
        indent=False,
    )

    add_section_heading(document, "", "References")
    # Remove the leading period used for numbered sections from the references heading.
    heading = document.paragraphs[-1]
    heading.runs[0].text = "REFERENCES"
    add_reference(
        document,
        1,
        "K. E. Ong, S. Retta, R. Srinivasan, S. Tan, and J. Liu, “CattleEyeView: A Multi-task Top-down View Cattle Dataset for Smarter Precision Livestock Farming,” arXiv:2312.08764, 2023.",
    )
    add_reference(
        document,
        2,
        "Ultralytics, “Model Training with Ultralytics YOLO,” Ultralytics YOLO Documentation. [Online]. Available: https://docs.ultralytics.com/modes/train/. Accessed: Jul. 24, 2026.",
    )
    add_reference(
        document,
        3,
        "PyTorch, “Transfer Learning for Computer Vision Tutorial,” PyTorch Tutorials. [Online]. Available: https://docs.pytorch.org/tutorials/beginner/transfer_learning_tutorial.html. Accessed: Jul. 24, 2026.",
    )
    return document


def main() -> None:
    OUTPUTS[0].parent.mkdir(parents=True, exist_ok=True)
    document = build_report()
    for output in OUTPUTS:
        document.save(output)
        with ZipFile(output) as package:
            if package.testzip() is not None:
                raise RuntimeError(f"The generated Word package failed its integrity check: {output}")
        verified = Document(output)
        column_counts = []
        for section in verified.sections:
            columns = section._sectPr.find(qn("w:cols"))
            column_counts.append(columns.get(qn("w:num"), "1") if columns is not None else "1")
        print(
            f"{output}\n"
            f"Verified: {len(verified.paragraphs)} paragraphs, {len(verified.tables)} tables, "
            f"{len(verified.inline_shapes)} figures, columns {column_counts}"
        )


if __name__ == "__main__":
    main()
